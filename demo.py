import streamlit as st
import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import tempfile
import concurrent.futures
import time
from tenacity import retry, stop_after_attempt, wait_exponential
import pandas as pd
from docx import Document
import openpyxl
import io

# Load environment variables
load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Increase chunk size and overlap for larger documents
CHUNK_SIZE = 1000000
CHUNK_OVERLAP = 100000


class RateLimiter:
    def __init__(self, max_calls, period):
        self.max_calls = max_calls
        self.period = period
        self.calls = []

    def __call__(self, f):
        def wrapper(*args, **kwargs):
            now = time.time()
            self.calls = [c for c in self.calls if c > now - self.period]
            if len(self.calls) >= self.max_calls:
                raise Exception("Rate limit exceeded. Please try again later.")
            self.calls.append(now)
            return f(*args, **kwargs)

        return wrapper


@st.cache_resource
def get_pdf_text(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def get_docx_text(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text


def get_excel_text(excel_file):
    df_dict = pd.read_excel(excel_file, sheet_name=None)
    text = ""
    for sheet_name, df in df_dict.items():
        text += f"\nSheet: {sheet_name}\n"
        # Convert headers to string
        headers = [str(col) for col in df.columns]
        text += "Headers: " + ", ".join(headers) + "\n"
        # Convert all data to string and join
        for index, row in df.iterrows():
            text += "Row " + str(index) + ": " + " | ".join([str(val) for val in row.values]) + "\n"
    return text


def get_txt_text(txt_file):
    return txt_file.getvalue().decode('utf-8')


def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = text_splitter.split_text(text)
    return chunks


@st.cache_resource
def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    return vector_store


def get_conversational_chain():
    prompt_template = """
    You are an AI assistant specializing in analyzing various document types including research papers, spreadsheets, and general documents. Answer the question based on the following guidelines:

    1. If the answer is in the provided context, use that information.
    2. For general knowledge questions not in the context, use your built-in knowledge.
    3. For document-specific questions, analyze the context and provide informed responses.
    4. For programming questions, provide explanations or code snippets as needed.
    5. For research paper analysis:
       - Summarize key findings and methodologies
       - Explain complex scientific concepts
       - Identify the main contributions
       - Suggest related papers or further research areas
    6. For spreadsheet data:
       - Identify patterns and trends
       - Summarize key metrics
       - Provide data-driven insights
    7. Always indicate the basis of your answer (context, general knowledge, analysis, or specific document type expertise).

    Context: {context}

    Question: {question}

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain


def is_research_paper(text):
    keywords = ['abstract', 'introduction', 'methodology', 'results', 'conclusion', 'references']
    return sum(1 for keyword in keywords if keyword in text.lower()) >= 4


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
@RateLimiter(max_calls=10, period=60)  # 10 calls per minute
def user_input_with_retry(user_question, vector_store):
    docs = vector_store.similarity_search(user_question)
    chain = get_conversational_chain()
    response = chain(
        {"input_documents": docs, "question": user_question},
        return_only_outputs=True
    )

    if "```" in response["output_text"]:
        formatted_response = format_code_in_response(response["output_text"])
        return formatted_response

    return response["output_text"]


def format_code_in_response(response):
    parts = response.split("```")
    formatted_response = parts[0]

    for i in range(1, len(parts), 2):
        code = parts[i].strip()
        formatted_response += f"\n\n```\n{code}\n```\n\n"
        if i + 1 < len(parts):
            formatted_response += parts[i + 1]

    return formatted_response


def is_programming_question(question):
    programming_keywords = ['code', 'program', 'function', 'algorithm', 'syntax', 'debug']
    return any(keyword in question.lower() for keyword in programming_keywords)


def process_document(file):
    # Create a temporary file to handle the uploaded file
    with tempfile.NamedTemporaryFile(delete=False, suffix="." + file.name.split('.')[-1]) as tmp_file:
        tmp_file.write(file.getvalue())
        tmp_file_path = tmp_file.name

    try:
        file_extension = file.name.split('.')[-1].lower()
        if file_extension == 'pdf':
            text = get_pdf_text(tmp_file_path)
        elif file_extension == 'docx':
            text = get_docx_text(tmp_file_path)
        elif file_extension in ['xlsx', 'xls']:
            text = get_excel_text(tmp_file_path)
        elif file_extension == 'txt':
            text = get_txt_text(file)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

        # Process the text based on document type
        if file_extension == 'pdf' and is_research_paper(text):
            sections = extract_research_paper_text(tmp_file_path)
            text = "\n\n".join([f"{section.upper()}:\n{content}" for section, content in sections.items()])

        return text

    finally:
        # Clean up the temporary file
        os.unlink(tmp_file_path)


def extract_research_paper_text(pdf_file):
    text = get_pdf_text(pdf_file)
    sections = identify_paper_sections(text)
    return sections


def identify_paper_sections(text):
    sections = {
        "abstract": "",
        "introduction": "",
        "methodology": "",
        "results": "",
        "discussion": "",
        "conclusion": "",
        "references": ""
    }
    current_section = ""
    for line in text.split('\n'):
        lower_line = line.lower()
        if "abstract" in lower_line:
            current_section = "abstract"
        elif "introduction" in lower_line:
            current_section = "introduction"
        elif "method" in lower_line or "methodology" in lower_line:
            current_section = "methodology"
        elif "result" in lower_line:
            current_section = "results"
        elif "discussion" in lower_line:
            current_section = "discussion"
        elif "conclusion" in lower_line:
            current_section = "conclusion"
        elif "reference" in lower_line:
            current_section = "references"

        if current_section:
            sections[current_section] += line + "\n"

    return sections


def get_general_knowledge_answer(question):
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
    response = model.predict(f"Answer this general knowledge question: {question}")
    return response


def main():
    st.set_page_config(page_title="AI Amatya", page_icon=":books:", layout="wide")
    st.header("AI Amatya 💬📚")

    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'user_name' not in st.session_state:
        st.session_state.user_name = ""
    if 'vector_store' not in st.session_state:
        st.session_state.vector_store = None

    with st.sidebar:
        st.image("logo.png", width=150)
        st.title("Menu")
        user_name = st.text_input("Your Name", value=st.session_state.user_name)
        if user_name != st.session_state.user_name:
            st.session_state.user_name = user_name

        st.title("File Upload")
        uploaded_files = st.file_uploader(
            "Upload your Documents",
            accept_multiple_files=True,
            type=["pdf", "docx", "xlsx", "xls", "txt"]
        )

        if st.button("Process Documents"):
            if uploaded_files:
                with st.spinner("Processing documents... This may take a while for large files."):
                    with concurrent.futures.ThreadPoolExecutor() as executor:
                        texts = list(executor.map(process_document, uploaded_files))

                    raw_text = "\n".join(texts)
                    text_chunks = get_text_chunks(raw_text)
                    st.session_state.vector_store = get_vector_store(text_chunks)
                    st.success("Documents processed successfully!")
            else:
                st.error("Please upload documents before processing.")

        if st.checkbox("Dark Mode"):
            st.markdown("""
                <style>
                    .stApp {
                        background-color: #2b2b2b;
                        color: white;
                    }
                </style>
            """, unsafe_allow_html=True)

        if st.session_state.vector_store is not None:
            st.sidebar.subheader("Document Analysis")
            if st.sidebar.button("Summarize Key Points"):
                summary = user_input_with_retry("Summarize the key points of the uploaded documents",
                                                st.session_state.vector_store)
                st.sidebar.write(summary)

            if st.sidebar.button("Identify Main Topics"):
                topics = user_input_with_retry("Identify the main topics discussed in the uploaded documents",
                                               st.session_state.vector_store)
                st.sidebar.write(topics)

            if st.sidebar.button("Suggest Further Reading"):
                suggestions = user_input_with_retry(
                    "Suggest related topics or areas for further reading based on the uploaded documents",
                    st.session_state.vector_store)
                st.sidebar.write(suggestions)

    st.subheader("Chat")
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.write(f"{message['content']}")

    user_question = st.chat_input("Ask a question about your documents or any other topic")
    if user_question:
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        with st.chat_message("user"):
            st.write(user_question)

        with st.spinner("Thinking...."):
            try:
                if st.session_state.vector_store is None and not is_programming_question(user_question):
                    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
                    response = model.predict(user_question)
                else:
                    response = user_input_with_retry(user_question, st.session_state.vector_store)

                st.session_state.chat_history.append({"role": "assistant", "content": response})
                with st.chat_message("assistant"):
                    st.markdown(response)
            except Exception as e:
                st.error(f"An error occurred: {str(e)}. Please try again later.")

    if st.button("Clear Chat"):
        st.session_state.chat_history = []
        st.experimental_rerun()

    st.markdown("---")
    st.markdown("Created with ❤️ by InnovAIt-ON")


if __name__ == "__main__":
    main()